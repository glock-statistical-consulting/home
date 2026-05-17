import os
import sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_theme import (
    MIDNIGHT, IVORY, STEEL, ORANGE, FONT_NAME,
    set_cell_shading, add_cell_paragraph, add_header, add_footer, page_setup_portrait
)
from data_einfuehrungen import (
    r_bachelor, r_master, python_bachelor, python_master, spss_bachelor, spss_master
)


def create_einfuehrung(title, subtitle, level, sprache, code_blocks, path):
    doc = Document()
    page_setup_portrait(doc)
    add_header(doc, topic="Einf\u00fchrung " + sprache, level=level)
    add_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT_NAME

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\u2501" * 80)
    run.font.size = Pt(3)
    run.font.color.rgb = ORANGE

    for idx, (section_title, steps) in enumerate(code_blocks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(str(idx + 1) + ".  " + section_title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

        for step_title, code, explanation, expected_output in steps:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(step_title)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = FONT_NAME

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(explanation)
            run.font.size = Pt(9.5)
            run.font.name = FONT_NAME

            for line in code.split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(line)
                run.font.size = Pt(8.5)
                run.font.name = "Consolas"
                run.font.color.rgb = MIDNIGHT

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run("Erwarteter Output:")
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.name = FONT_NAME
            run.font.color.rgb = RGBColor(0x4F, 0x6D, 0x8A)

            for line in expected_output.split("\n"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(line)
                run.font.size = Pt(8.5)
                run.font.name = "Consolas"
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    import shutil, time
    tmpdir = os.path.join(os.path.dirname(path), "_tmp")
    os.makedirs(tmpdir, exist_ok=True)
    tmp = os.path.join(tmpdir, os.path.basename(path))
    doc.save(tmp)
    for attempt in range(10):
        try:
            if os.path.exists(path):
                os.remove(path)
            shutil.move(tmp, path)
            break
        except (PermissionError, OSError):
            time.sleep(1)
    print("Erstellt: " + os.path.basename(path))


base = r"C:\Users\Kevalon\OneDrive\Documents\GlockConsulting\local\downloads\05_introductions"

create_einfuehrung(
    "Einf\u00fchrung in R \u2014 Standard Level",
    "Grundlagen \u00b7 Vektoren \u00b7 Data Frames \u00b7 t-Test \u00b7 Lineare Regression",
    "Standard", "R", r_bachelor,
    os.path.join(base, "introduction_R_standard.docx")
)

create_einfuehrung(
    "Einf\u00fchrung in R \u2014 Advanced Level",
    "dplyr \u00b7 Funktionen \u00b7 Multiple Regression \u00b7 ANOVA \u00b7 ggplot2",
    "Advanced", "R", r_master,
    os.path.join(base, "introduction_R_advanced.docx")
)

create_einfuehrung(
    "Einf\u00fchrung in Python \u2014 Standard Level",
    "NumPy \u00b7 Pandas \u00b7 t-Test \u00b7 Korrelation",
    "Standard", "Python", python_bachelor,
    os.path.join(base, "introduction_python_standard.docx")
)

create_einfuehrung(
    "Einf\u00fchrung in Python \u2014 Advanced Level",
    "groupby \u00b7 statsmodels \u00b7 ANOVA \u00b7 Seaborn",
    "Advanced", "Python", python_master,
    os.path.join(base, "introduction_python_advanced.docx")
)

create_einfuehrung(
    "Einf\u00fchrung in SPSS \u2014 Standard Level",
    "Frequenzen \u00b7 Recode \u00b7 t-Test \u00b7 Korrelation",
    "Standard", "SPSS", spss_bachelor,
    os.path.join(base, "introduction_spss_standard.docx")
)

create_einfuehrung(
    "Einf\u00fchrung in SPSS \u2014 Advanced Level",
    "UNIANOVA \u00b7 Logistische Regression \u00b7 Reliabilit\u00e4tsanalyse",
    "Advanced", "SPSS", spss_master,
    os.path.join(base, "introduction_spss_advanced.docx")
)

print("\nAlle Einf\u00fchrungsdokumente erfolgreich erstellt.")
