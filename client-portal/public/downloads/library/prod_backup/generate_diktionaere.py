import os, sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_theme import (
    MIDNIGHT, IVORY, STEEL, ORANGE, FONT_NAME, set_cell_shading, add_cell_paragraph,
    add_header, add_footer, page_setup_portrait, set_row_keep_together
)
from data_diktionaere import abi, bachelor, master, phd


def create_dictionary(title, level, data, path):
    doc = Document()
    page_setup_portrait(doc)
    add_header(doc, topic=title, level=level)
    add_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True; run.font.size = Pt(16); run.font.name = FONT_NAME

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deutsch-Englisches Glossar statistischer Fachbegriffe")
    run.font.size = Pt(10); run.font.italic = True; run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    category_order = []
    cat_items = {}
    for de, en, def_de, def_en, cat in data:
        if cat not in cat_items:
            category_order.append(cat)
            cat_items[cat] = []
        cat_items[cat].append((de, en, def_de, def_en))

    for cat_idx, cat in enumerate(category_order):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10) if cat_idx > 0 else Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(cat)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_NAME
        run.font.color.rgb = ORANGE

        items = cat_items[cat]
        table = doc.add_table(rows=len(items) + 1, cols=4)
        table.autofit = True
        table.style = 'Table Grid'

        headers = ["#", "Deutsch", "English", "Definition"]
        for i, h in enumerate(headers):
            set_cell_shading(table.cell(0, i), MIDNIGHT)
            add_cell_paragraph(table.cell(0, i), h, bold=True, size=9, color=IVORY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for row_idx, (de, en, def_de, def_en) in enumerate(items):
            row = table.rows[row_idx + 1]
            set_row_keep_together(row)
            add_cell_paragraph(row.cells[0], str(row_idx + 1), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_cell_paragraph(row.cells[1], de, bold=True, size=8)
            add_cell_paragraph(row.cells[2], en, bold=True, size=8)
            add_cell_paragraph(row.cells[3], def_de, size=8)

        doc.add_paragraph()

    doc.save(path)
    print(f"Erstellt: {os.path.basename(path)}")


base = r"C:\Users\Kevalon\OneDrive\Documents\GlockConsulting\local\downloads\04_dictionaries"

create_dictionary("Diktion\u00e4r \u2014 Basic", "Basic", abi, os.path.join(base, "dictionary_basic.docx"))
create_dictionary("Diktion\u00e4r \u2014 Standard", "Standard", bachelor, os.path.join(base, "dictionary_standard.docx"))
create_dictionary("Diktion\u00e4r \u2014 Advanced", "Advanced", master, os.path.join(base, "dictionary_advanced.docx"))
create_dictionary("Diktion\u00e4r \u2014 All-Access", "All-Access", phd, os.path.join(base, "dictionary_all_access.docx"))

print("\nAlle vier Diktionäre erfolgreich erstellt.")
