import os
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

MIDNIGHT = RGBColor(0x14, 0x21, 0x3D)
IVORY = RGBColor(0xFA, 0xFA, 0xF8)
STEEL = RGBColor(0x4F, 0x6D, 0x8A)
ORANGE = RGBColor(0xFF, 0x66, 0x00)
FONT_NAME = "Calibri"
FONT_SIZE = 9.5

def page_setup_portrait(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

def page_setup_landscape(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)

def add_header(doc, topic="", level=""):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    logo_path = os.path.join(os.path.dirname(__file__), "logo-compact.png")
    if os.path.exists(logo_path):
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2.8))

    run = p.add_run("  Kevin Glock Statistical Consulting Services")
    run.font.size = Pt(7.5)
    run.font.name = FONT_NAME
    run.font.color.rgb = STEEL
    run.font.italic = True

    pr = p.add_run()
    pr.font.size = Pt(7.5)
    pr.add_tab()

    tp = "  |  " + topic
    run3 = p.add_run(tp)
    run3.font.size = Pt(7.5)
    run3.font.name = FONT_NAME
    run3.font.color.rgb = MIDNIGHT

def add_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run = p.add_run("Kevin Glock Statistical Consulting Services  |  GSC Services")
    run.font.size = Pt(7)
    run.font.name = FONT_NAME
    run.font.color.rgb = STEEL
    run.font.italic = True

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_cell_paragraph(cell, text, bold=False, size=9, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    run.bold = bold
    if color:
        run.font.color.rgb = color

def set_row_keep_together(row):
    for cell in row.cells:
        tc = cell._tc
        p = tc.get_or_add_tcPr()
        cant_split = OxmlElement('w:cantSplit')
        cant_split.set(qn('w:val'), 'true')
        p.append(cant_split)

def set_keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    keep.set(qn('w:val'), 'true')
    pPr.append(keep)
