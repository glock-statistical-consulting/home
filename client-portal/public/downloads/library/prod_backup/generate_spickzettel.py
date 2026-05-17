import os, sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx_theme import (
    MIDNIGHT, STEEL, ORANGE, FONT_NAME, set_cell_shading, add_header, add_footer, page_setup_landscape
)

def set_cell(cell, text, bold=False, size=8, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT_NAME
    r.bold = bold
    if color:
        r.font.color.rgb = color

def make_sheet(fname, title, subtitle, rows, col_widths):
    doc = Document()
    page_setup_landscape(doc)
    add_header(doc, topic="Formelsammlung / Spickzettel")
    add_footer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.name = FONT_NAME
    r.font.color.rgb = MIDNIGHT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(9); r.font.italic = True; r.font.name = FONT_NAME
    r.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("\u2501" * 100)
    r.font.size = Pt(3)
    r.font.color.rgb = ORANGE
    table = doc.add_table(rows=len(rows), cols=len(col_widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)
    for idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(idx, col_idx)
            if idx == 0:
                set_cell_shading(cell, MIDNIGHT)
                set_cell(cell, cell_data, bold=True, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if col_idx == 0:
                    set_cell(cell, cell_data, bold=True, size=8, color=RGBColor(0x14, 0x21, 0x3D))
                else:
                    set_cell(cell, cell_data, size=8)
    folder = r"C:\Users\Kevalon\OneDrive\Documents\GlockConsulting\local\downloads\03_cheatsheets"
    doc.save(os.path.join(folder, fname))
    print("  " + fname)

make_sheet("spickzettel_abi.docx",
    "Formelsammlung Statistik — Abitur / Basic Level",
    "Deskriptive Statistik, Wahrscheinlichkeit, Hypothesentests Grundlagen",
    [
        ["Thema", "Formel / Konzept", "Erklärung / Beispiel"],
        ["Mittelwert", "x̄ = (1/n) · Σxᵢ", "Arithmetisches Mittel. Beispiel: (2+4+6)/3 = 4"],
        ["Median", "Sortierte Liste, mittlerer Wert", "n=7: 4. Wert. n=6: Mitte der beiden mittleren"],
        ["Modalwert (Modus)", "Häufigster Wert", "In [1,1,2,3,3,3,4] ist 3 der Modus"],
        ["Spannweite", "R = max − min", "In [2,5,8,12]: R = 12−2 = 10"],
        ["Varianz", "s² = 1/(n−1) · Σ(xᵢ − x̄)²", "Mittlere quadratische Abweichung vom Mittelwert"],
        ["Standardabweichung", "s = √s²", "Wurzel der Varianz. Gleiche Einheit wie Daten"],
        ["Variationskoeffizient", "V = s / x̄ · 100", "Relative Streuung in %. Nur Ratioskala"],
        ["Standardisierung", "z = (x − x̄) / s", "z-Wert: Abweichung in SD-Einheiten. Mittelwert 0, SD 1"],
        ["Korrelation (r)", "r = cov(x,y) / (sₓ · sᵧ)", "Bereich: −1 bis +1. 0 = kein linearer Zusammenhang"],
        ["Anteilswert", "p̂ = k / n", "Anteil von k Erfolgen in n Versuchen"],
        ["Wahrscheinlichkeit", "P(A) = |A| / |Ω|", "Laplace: günstige / mögliche Ergebnisse"],
        ["Bedingte WS", "P(A|B) = P(A∩B) / P(B)", "Wahrscheinlichkeit von A, gegeben B"],
        ["Unabhängigkeit", "P(A∩B) = P(A) · P(B)", "Wenn A und B unabhängig sind"],
        ["Erwartungswert", "E(X) = Σ(xᵢ · pᵢ)", "Mittelwert der Zufallsvariable X"],
        ["Binomialverteilung", "P(X=k) = C(n,k) · pᵏ · (1−p)ⁿ⁻ᵏ", "n Versuche, Erfolgswahrscheinlichkeit p"],
        ["Normalverteilung", "φ(z) = 1/√(2π) · e^(−z²/2)", "Glockenkurve. Standardnormal N(0,1)"],
        ["Sigma-Regeln", "μ ± 1σ: 68%; ± 2σ: 95%; ± 3σ: 99,7%", "Faustregel für Normalverteilung"],
        ["KI (Mittelwert)", "KI = x̄ ± z · s / √n", "z.B. z=1,96 für 95%-KI"],
        ["Signifikanzniveau", "α = 0,05 (5%)", "Wahrscheinlichkeit für Fehler 1. Art"],
        ["p-Wert", "p < α → H₀ ablehnen", "Wahrscheinlichkeit des Ergebnisses unter H₀"],
        ["Fehler 1. Art", "H₀ abgelehnt, obwohl wahr", "Falsch-Positiv. Risiko = α"],
        ["Fehler 2. Art", "H₀ beibehalten, obwohl H₁ wahr", "Falsch-Negativ. Risiko = β"],
    ],
    [3.5, 8.0, 14.0]
)

make_sheet("spickzettel_bachelor.docx",
    "Formelsammlung Statistik — Bachelor / Standard Level",
    "Induktive Statistik, t-Test, ANOVA, Regression, Chi-Quadrat",
    [
        ["Thema", "Formel / Konzept", "Erklärung / Beispiel"],
        ["Einstichproben-t", "t = (x̄ − μ₀) / (s / √n)", "Vergleich Stichprobenmittel mit Population. df = n−1"],
        ["Unabhängiger t-Test", "t = (x̄₁ − x̄₂) / √(s²p · (1/n₁ + 1/n₂))", "Vergleich zweier unabhängiger Gruppen. s²p = gepoolte Varianz"],
        ["Verbundener t-Test", "t = d̄ / (sᵈ / √n)", "Differenz pro Person. df = n−1"],
        ["Cohens d", "d = (x̄₁ − x̄₂) / s_pooled", "Effektstärke: 0,2=klein, 0,5=mittel, 0,8=groß"],
        ["F-Statistik (ANOVA)", "F = MS_zwischen / MS_innerhalb", "Varianz zwischen / Varianz innerhalb Gruppen"],
        ["Post-hoc Bonferroni", "α_korrigiert = α / k", "k = Anzahl Vergleiche. Konservativ"],
        ["Tukey HSD", "HSD = q · √(MS_inner / n)", "Ehrlich signifikanter Unterschied"],
        ["Eta-Quadrat", "η² = SS_zwischen / SS_gesamt", "Anteil erklärter Varianz durch Gruppen"],
        ["Einfache Regression", "ŷ = b₀ + b₁·x", "b₁ = r · (sᵧ/sₓ); b₀ = ȳ − b₁·x̄"],
        ["Bestimmtheitsmaß", "R² = SS_Regression / SS_gesamt", "Anteil erklärter Varianz. 0 bis 1"],
        ["SE der Regression", "SE = √(Σeᵢ² / (n−2))", "Standardfehler der Residuen"],
        ["Korrelation", "r = Σ((x−x̄)(y−ȳ)) / ((n−1)·sₓ·sᵧ)", "Pearson-Korrelation. −1 ≤ r ≤ +1"],
        ["Chi-Quadrat", "χ² = Σ((O−E)² / E)", "O = beobachtet, E = erwartet. Unabhängigkeitstest"],
        ["Cramers V", "V = √(χ² / (n · (min(r,c)−1)))", "Effektstärke für χ²: 0,1 / 0,3 / 0,5"],
        ["Fishers exakter Test", "Hypergeometrische WS", "Für 2×2-Tabellen mit Zellen < 5"],
        ["KI Mittelwert", "KI = x̄ ± t_(n−1, α/2) · s / √n", "t-Wert aus Tabelle"],
        ["KI Anteil", "KI = p̂ ± z · √(p̂·(1−p̂)/n)", "z.B. z=1,96 für 95%"],
        ["Stichprobengröße", "n = (z² · p · (1−p)) / e²", "e = gewünschte Fehlermarge (z.B. 0,03)"],
        ["Standardfehler", "SE = s / √n", "Streuung der Stichprobenverteilung"],
        ["Effektstärke (χ²)", "w = √(χ² / N)", "Cohen: 0,1=klein, 0,3=mittel, 0,5=groß"],
    ],
    [3.5, 8.0, 14.0]
)

make_sheet("spickzettel_master.docx",
    "Formelsammlung Statistik — Master / Advanced Level",
    "Multiple Regression, SEM, Faktorenanalyse, Nichtparametrik",
    [
        ["Thema", "Formel / Konzept", "Erklärung / Beispiel"],
        ["Multiple Regression", "y = b₀ + b₁x₁ + ... + bₖxₖ + e", "Mehrere UVs gleichzeitig. Partielle Regressionskoeffizienten"],
        ["Korrigiertes R²", "R²_adj = 1 − (1−R²)·(n−1)/(n−k−1)", "Bestraft viele Prädiktoren. Modellvergleich"],
        ["VIF", "VIF = 1 / (1 − R²ⱼ)", "R²ⱼ aus Regression von Xⱼ auf andere. VIF > 10 = problematisch"],
        ["Partielle Korrelation", "r_xy.z = (r_xy − r_xz·r_yz) / √((1−r²_xz)·(1−r²_yz))", "Korrelation von x und y nach Auspartialisierung von z"],
        ["Logit (Log. Regr.)", "ln(p/(1−p)) = b₀ + b₁x₁ + ... + bₖxₖ", "Log-Odds als lineare Funktion der Prädiktoren"],
        ["Odds-Ratio", "OR = exp(b)", "Faktoränderung der Odds pro X-Einheit. OR > 1 = steigend"],
        ["Pseudo-R² (Nagelkerke)", "R²_N = (1 − (L₀/Lₘ)^(2/n)) / (1 − L₀^(2/n))", "Approximation erklärter Varianz"],
        ["AIC", "AIC = −2·ln(L) + 2·k", "k = Parameteranzahl. Niedriger AIC = besseres Modell"],
        ["BIC", "BIC = −2·ln(L) + k·ln(n)", "Strenger als AIC (bevorzugt einfachere Modelle)"],
        ["Partielles η²", "η²_p = SS_Effekt / (SS_Effekt + SS_Fehler)", "Effektstärke pro Effekt in ANOVA/ANCOVA"],
        ["Eigenwert (FA)", "EW = Σa²ᵢ für Faktor i", "aᵢ = Ladungen. Kaiser: EW > 1 behalten"],
        ["KMO", "KMO = Σr²ᵢⱼ / (Σr²ᵢⱼ + Σa²ᵢⱼ)", "KMO > 0,6 = akzeptabel. Stichprobeneignung"],
        ["Bartlett-Test", "χ² = −((n−1) − (2p+5)/6)·ln(|R|)", "Test auf nicht-diagonale Struktur. p < ,05 = geeignet"],
        ["CFI (CFA)", "CFI = 1 − max(χ²_m − df_m, 0) / max(χ²_m − df_m, χ²_b − df_b, 0)", "CFI > ,95 = guter Fit. Ziel- vs. Basismodell"],
        ["RMSEA", "RMSEA = √(max((χ²_m − df_m)/(N−1), 0) / df_m)", "RMSEA < ,06 = gut; < ,08 akzeptabel"],
        ["Mann-Whitney-U", "U = ΣR₁ − n₁·(n₁+1)/2", "Rangsumme Gruppe 1 minus Minimum"],
        ["Kruskal-Wallis-H", "H = 12/(N(N+1)) · Σ(R²ᵢ/nᵢ) − 3(N+1)", "Rangbasierte ANOVA ~ χ² mit k−1 df"],
        ["Friedman", "χ²_F = 12n/(k(k+1)) · Σ(R²ⱼ) − 3n(k+1)", "Rangbasierte RM-ANOVA. Rⱼ = Rangsumme pro Zeit"],
    ],
    [3.5, 8.0, 14.0]
)

make_sheet("spickzettel_phd.docx",
    "Formelsammlung Statistik — PhD / All-Access Level",
    "Bayes, Mehrebenen, SEM, ML, Meta-Analyse",
    [
        ["Thema", "Formel / Konzept", "Erklärung / Beispiel"],
        ["Bayes-Theorem", "P(θ|D) = P(D|θ)·P(θ) / P(D)", "Posterior = Likelihood · Prior / Evidence"],
        ["Bayes-Faktor", "BF₁₀ = P(D|H₁) / P(D|H₀)", "BF > 3: moderat; BF > 10: starke Evidenz"],
        ["MCMC", "P(θ|D) ≈ 1/m · Σg(θ^(t))", "Posterior-Mean via MCMC. R-hat < 1,01"],
        ["ICC (Mehrebenen)", "ICC = σ²_zwischen / (σ²_zwischen + σ²_innerhalb)", "Anteil Level-2-Varianz. ICC > ,05 ⇒ HLM"],
        ["Mehrebenenmodell", "yᵢⱼ = γ₀₀ + γ₁₀·xᵢⱼ + u₀ⱼ + u₁ⱼ·xᵢⱼ + eᵢⱼ", "Fixed: γ; Random: u; Level-1: i; Level-2: j"],
        ["Cross-Level-Interaktion", "γ₁₁ = Effekt von wⱼ (L2) auf b₁ⱼ (L1-Steigung)", "Moderiert eine L2-Variable einen L1-Effekt?"],
        ["SEM – CFI", "CFI = 1 − max(χ²_m − df_m, 0) / max(χ²_m − df_m, χ²_b − df_b, 0)", "> ,95: gut; > ,90: akzeptabel"],
        ["SEM – RMSEA", "RMSEA = √(max(χ²_m − df_m, 0) / (df_m · (N−1)))", "< ,06: gut; ,06−,08: akzeptabel; > ,10: schlecht"],
        ["Mediation (PROCESS)", "a·b = indirekter Effekt; c' = direkter Effekt", "X→M (a); M→Y (b); a·b Boot-KI ≠ 0 = sign. Mediation"],
        ["Moderation", "y = b₀ + b₁x + b₂z + b₃x·z + e", "b₃ = Interaktion. Simple Slopes bei Mean ± 1 SD von z"],
        ["Poweranalyse", "n = f(α, β, d)", "Power > ,80. G*Power. Effektstärke aus Vorforschung"],
        ["Meta (Fixeffekt)", "θ_gepoolt = Σ(wᵢ·θᵢ) / Σwᵢ", "wᵢ = 1/vᵢ (inverse Varianz). Ein wahrer Effekt"],
        ["Meta (Zufallseffekt)", "w*ᵢ = 1 / (vᵢ + τ²)", "τ² = Zwischenstudien-Varianz. Effekt variiert"],
        ["I²", "I² = (Q − df) / Q · 100%", "Heterogenitätsanteil. 25/50/75% = niedrig/mittel/hoch"],
        ["Funnel-Plot", "Effekt (x) vs. SE (y)", "Asymmetrie = Publikationsbias. Egger-Test p < ,10"],
        ["Präzision-Recall", "P = TP/(TP+FP); R = TP/(TP+FN)", "Präzision: Korrektheit. Recall: Vollständigkeit"],
        ["F1-Score", "F1 = 2 · P · R / (P + R)", "Harmonisches Mittel von Präzision und Recall"],
        ["AUC-ROC", "AUC = ∫(TPR(FPR)) d(FPR)", "AUC > ,70 akzeptabel; > ,80 gut; > ,90 exzellent"],
        ["Bias-Varianz", "E[(y − f̂(x))²] = Bias² + Var + σ²", "Bias: systematischer Fehler; Var: Streuung. Tradeoff"],
    ],
    [3.5, 8.0, 14.0]
)

print("\nSpickzettel erfolgreich erstellt.")
