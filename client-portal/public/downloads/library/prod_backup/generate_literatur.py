import os, sys, shutil, time
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_theme import FONT_NAME, add_header, add_footer, page_setup_portrait

sections = [
    ("Grundlagen & Methoden", [
        "Bortz, J. & Schuster, C. (2010). *Statistik f\u00fcr Human- und Sozialwissenschaftler* (7. Aufl.). Springer.",
        "Casella, G. & Berger, R. L. (2002). *Statistical Inference* (2nd ed.). Duxbury.",
        "Wasserman, L. (2004). *All of Statistics: A Concise Course in Statistical Inference*. Springer.",
        "Kruschke, J. K. (2015). *Doing Bayesian Data Analysis: A Tutorial with R, JAGS, and Stan* (2nd ed.). Academic Press.",
        "McElreath, R. (2020). *Statistical Rethinking: A Bayesian Course with Examples in R and Stan* (2nd ed.). CRC Press.",
        "Embretson, S. E. & Reise, S. P. (2000). *Item Response Theory for Psychologists*. Erlbaum.",
        "Reise, S. P. & Waller, N. G. (2009). Item response theory and clinical measurement. *Annual Review of Clinical Psychology, 5*, 27\u201348. https://doi.org/10.1146/annurev.clinpsy.032408.153553",
        "Schnell, R., Hill, P. B. & Esser, E. (2018). *Methoden der empirischen Sozialforschung* (11. Aufl.). De Gruyter.",
        "Bortz, J. & D\u00f6ring, N. (2006). *Forschungsmethoden und Evaluation f\u00fcr Human- und Sozialwissenschaftler* (4. Aufl.). Springer.",
        "Diekmann, A. (2018). *Empirische Sozialforschung: Grundlagen, Methoden, Anwendungen* (12. Aufl.). Rowohlt.",
        "Fahrmeir, L., Kneib, T., Lang, S. & Marx, B. (2013). *Regression: Models, Methods and Applications*. Springer.",
        "Cohen, J. (1988). *Statistical Power Analysis for the Behavioral Sciences* (2nd ed.). Erlbaum.",
        "Shadish, W. R., Cook, T. D. & Campbell, D. T. (2002). *Experimental and Quasi-Experimental Designs for Generalized Causal Inference*. Houghton Mifflin.",
    ]),
    ("Angewandte Verfahren", [
        "Backhaus, K., Erichson, B., Gensler, S., Weiber, R. & Weiber, T. (2021). *Multivariate Analysemethoden: Eine anwendungsorientierte Einf\u00fchrung* (16. Aufl.). Springer Gabler.",
        "Hair, J. F., Black, W. C., Babin, B. J. & Anderson, R. E. (2019). *Multivariate Data Analysis* (8th ed.). Cengage.",
        "Tabachnick, B. G. & Fidell, L. S. (2019). *Using Multivariate Statistics* (7th ed.). Pearson.",
        "Field, A. (2018). *Discovering Statistics Using IBM SPSS Statistics* (5th ed.). Sage.",
        "Field, A. (2018). *Statistik mit R* (2. Aufl.). Springer. [\u00dcbersetzung von *Discovering Statistics Using R*]",
        "Field, A., Miles, J. & Field, Z. (2012). *Discovering Statistics Using R*. Sage.",
        "B\u00fchner, M. & Ziegler, M. (2017). *Statistik f\u00fcr Psychologen und Sozialwissenschaftler* (2. Aufl.). Pearson.",
        "Eid, M., Gollwitzer, M. & Schmitt, M. (2017). *Statistik und Forschungsmethoden* (5. Aufl.). Beltz.",
    ]),
    ("Vertiefte Verfahren", [
        "Muth\u00e9n, L. K. & Muth\u00e9n, B. O. (1998\u20132017). *Mplus User\u2019s Guide* (8th ed.). Muth\u00e9n & Muth\u00e9n.",
        "Kline, R. B. (2016). *Principles and Practice of Structural Equation Modeling* (4th ed.). Guilford.",
        "Brown, T. A. (2015). *Confirmatory Factor Analysis for Applied Research* (2nd ed.). Guilford.",
        "Hox, J. J., Moerbeek, M. & van de Schoot, R. (2018). *Multilevel Analysis: Techniques and Applications* (3rd ed.). Routledge.",
        "Gelman, A. & Hill, J. (2006). *Data Analysis Using Regression and Multilevel/Hierarchical Models*. Cambridge University Press.",
        "Agresti, A. (2013). *Categorical Data Analysis* (3rd ed.). Wiley.",
        "Hosmer, D. W., Lemeshow, S. & Sturdivant, R. X. (2013). *Applied Logistic Regression* (3rd ed.). Wiley.",
        "Box, G. E. P., Jenkins, G. M., Reinsel, G. C. & Ljung, G. M. (2015). *Time Series Analysis: Forecasting and Control* (5th ed.). Wiley.",
    ]),
    ("Software & Programmierung", [
        "R Core Team (2024). *R: A Language and Environment for Statistical Computing*. R Foundation for Statistical Computing. https://www.R-project.org/",
        "Wickham, H. & Grolemund, G. (2017). *R for Data Science*. O\u2019Reilly.",
        "Wickham, H. (2016). *ggplot2: Elegant Graphics for Data Analysis* (2nd ed.). Springer.",
        "McKinney, W. (2022). *Python for Data Analysis* (3rd ed.). O\u2019Reilly.",
        "VanderPlas, J. (2016). *Python Data Science Handbook*. O\u2019Reilly.",
        "Harris, C. R. et al. (2020). Array programming with NumPy. *Nature, 585*, 357\u2013362. https://doi.org/10.1038/s41586-020-2649-2",
    ]),
    ("Spezielle Themen", [
        "Mayring, P. (2015). *Qualitative Inhaltsanalyse: Grundlagen und Techniken* (12. Aufl.). Beltz.",
        "Corbin, J. & Strauss, A. (2015). *Basics of Qualitative Research: Techniques and Procedures for Developing Grounded Theory* (4th ed.). Sage.",
        "Keller, R. (2011). *Diskursforschung: Eine Einf\u00fchrung f\u00fcr SozialwissenschaftlerInnen* (4. Aufl.). VS Verlag.",
        "Flick, U. (2019). *Qualitative Sozialforschung: Eine Einf\u00fchrung* (9. Aufl.). Rowohlt.",
        "Creswell, J. W. & Creswell, J. D. (2018). *Research Design: Qualitative, Quantitative, and Mixed Methods Approaches* (5th ed.). Sage.",
        "Higgins, J. P. T. & Green, S. (Eds.) (2011). *Cochrane Handbook for Systematic Reviews of Interventions* (Version 5.1.0). https://training.cochrane.org/handbook",
        "Borenstein, M., Hedges, L. V., Higgins, J. P. T. & Rothstein, H. R. (2009). *Introduction to Meta-Analysis*. Wiley.",
    ]),
]

base = r"C:\Users\Kevalon\OneDrive\Documents\GlockConsulting\local\downloads"

doc = Document()
page_setup_portrait(doc)
add_header(doc, topic="Literaturverzeichnis Statistik & Methoden")
add_footer(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Literaturverzeichnis")
run.bold = True
run.font.size = Pt(16)
run.font.name = FONT_NAME

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Statistische Methoden \u00b7 Forschungsdesign \u00b7 Datenanalyse")
run.font.size = Pt(10)
run.font.name = FONT_NAME
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

overall = 0
for section_title, items in sections:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(section_title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

    for i, text in enumerate(items):
        overall += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        run = p.add_run(f"[{overall}]  ")
        run.font.size = Pt(9)
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
        parts = text.split("*")
        for i, part in enumerate(parts):
            r = p.add_run(part)
            r.font.size = Pt(9)
            r.font.name = FONT_NAME
            if i % 2 == 1:
                r.italic = True

path = os.path.join(base, "07_references", "literaturverzeichnis_statistik_methoden.docx")
tmpdir = os.path.join(base, "_tmp")
os.makedirs(tmpdir, exist_ok=True)
tmp = os.path.join(tmpdir, "literaturverzeichnis.docx")
doc.save(tmp)
for attempt in range(10):
    try:
        if os.path.exists(path):
            os.remove(path)
        shutil.move(tmp, path)
        break
    except (PermissionError, OSError):
        time.sleep(1)

print(f"Literaturverzeichnis aktualisiert: {len(sections)} Kategorien, {overall} Eintr\u00e4ge.")
