import os
import sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_theme import (
    MIDNIGHT, IVORY, STEEL, ORANGE, FONT_NAME, FONT_SIZE,
    set_cell_shading, add_cell_paragraph, add_header, add_footer, page_setup_portrait,
    set_keep_with_next
)
from data_klausur import bachelor_fragen, master_fragen


def add_frage(doc, num, text, typ, punkte):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}.  [{typ}]  ({punkte} Pkt.)  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = FONT_NAME

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run3 = p2.add_run("\u2501" * 80)
    run3.font.size = Pt(3)
    run3.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)


def create_klausur(title, subtitle, level, fragen, path):
    doc = Document()
    page_setup_portrait(doc)
    add_header(doc, topic="Klausurfragen Statistik", level=level)
    add_footer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT_NAME

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    gesamt_punkte = sum(p for *_, p in fragen)

    hinweise = [
        "Trage auf dem Deckblatt deinen Namen und deine Matrikelnummer ein.",
        f"Bearbeitungszeit: 90 Minuten.  Gesamtpunktzahl: {gesamt_punkte} Pkt.",
        "Erlaubte Hilfsmittel: Formelsammlung, Taschenrechner, Stift.",
        "Teil 1 (MC): Ankreuzen. Pro Frage ist genau eine Antwort richtig.",
        "Teil 2 (Offene Fragen): Kurze präzise Antworten.",
        "Teil 3 (Anwendungsaufgaben): Lösungswege dokumentieren. Ergebnisse klar kennzeichnen.",
        "Die Lösungen befinden sich im Anhang zur Selbstkontrolle."
    ]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("\U0001f4cb  HINWEISE")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

    for h in hinweise:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"\u2022  {h}")
        run.font.size = Pt(9)
        run.font.name = FONT_NAME

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2501" * 80)
    run.font.size = Pt(4)
    run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

    aktueller_teil = None
    for num, text, loesung, typ, teil, punkte in fragen:
        if teil != aktueller_teil:
            aktueller_teil = teil
            doc.add_paragraph()
            p = doc.add_paragraph()
            set_keep_with_next(p)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(teil)
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = FONT_NAME
            run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)
            run2 = p.add_run("\n" + "\u2501" * 50)
            run2.font.size = Pt(3)
            run2.font.color.rgb = RGBColor(0x4F, 0x6D, 0x8A)

        add_frage(doc, num, text, typ, punkte)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("LÖSUNGEN")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2014 Zur Selbstkontrolle \u2014")
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    aktueller_teil = None
    for num, text, loesung, typ, teil, punkte in fragen:
        if teil != aktueller_teil:
            aktueller_teil = teil
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(teil)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = FONT_NAME
            run.font.color.rgb = RGBColor(0x14, 0x21, 0x3D)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{num}.  ")
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

        run2 = p.add_run(loesung)
        run2.font.size = Pt(9.5)
        run2.font.name = FONT_NAME
        run2.font.color.rgb = RGBColor(0x4F, 0x6D, 0x8A)

    doc.save(path)
    print(f"Erstellt: {os.path.basename(path)}")


base = r"C:\Users\Kevalon\OneDrive\Documents\GlockConsulting\local\downloads\06_exam_questions"

create_klausur(
    "Klausurfragen \u2014 Statistik (Standard Level)",
    "Konstruierte Pr\u00fcfung \u00b7 Deskriptive & Induktive Statistik \u00b7 Inkl. L\u00f6sungshinweisen",
    "Standard",
    bachelor_fragen,
    os.path.join(base, "exam_questions_standard.docx")
)

create_klausur(
    "Klausurfragen \u2014 Statistik (Advanced Level)",
    "Konstruierte Pr\u00fcfung \u00b7 Fortgeschrittene Verfahren \u00b7 Inkl. L\u00f6sungshinweisen",
    "Advanced",
    master_fragen,
    os.path.join(base, "exam_questions_advanced.docx")
)

print("\nBeide Klausurfragen-Paper erfolgreich erstellt.")
