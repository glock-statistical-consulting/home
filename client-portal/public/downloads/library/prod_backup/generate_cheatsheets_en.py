import os, sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx_theme import (
    MIDNIGHT, STEEL, ORANGE, FONT_NAME, set_cell_shading, add_header, add_footer, page_setup_landscape
)

def set_cell(cell, text, bold=False, size=8, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT_NAME
    r.bold = bold
    if color:
        r.font.color.rgb = color

def make_sheet(fname, title, subtitle, rows, col_widths):
    doc = Document()
    page_setup_landscape(doc)
    add_header(doc, topic="Formula Sheet / Cheat Sheet")
    add_footer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(16); r.font.name = FONT_NAME
    r.font.color.rgb = MIDNIGHT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(9); r.font.italic = True; r.font.name = FONT_NAME
    r.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("\u2501" * 100)
    r.font.size = Pt(3)
    r.font.color.rgb = ORANGE
    table = doc.add_table(rows=len(rows), cols=len(col_widths))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)
    for idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(idx, col_idx)
            if idx == 0:
                set_cell_shading(cell, MIDNIGHT)
                set_cell(cell, cell_data, bold=True, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if col_idx == 0:
                    set_cell(cell, cell_data, bold=True, size=8, color=RGBColor(0x14, 0x21, 0x3D))
                else:
                    set_cell(cell, cell_data, size=8)
    folder = r"C:\Users\Kevalon\OneDrive\Documents\GlockConsulting\local\downloads\03_cheatsheets"
    doc.save(os.path.join(folder, fname))
    print("  " + fname)

make_sheet("cheatsheet_high_school.docx",
    "Statistics Formula Sheet — High School / Basic Level",
    "Descriptive Statistics, Probability, Hypothesis Testing Basics",
    [
        ["Topic", "Formula / Concept", "Explanation / Example"],
        ["Mean", "x_bar = (1/n) * Sum(x_i)", "Arithmetic mean. Example: (2+4+6)/3 = 4"],
        ["Median", "Sorted order, middle value", "n=7: 4th value. n=6: mean of middle two"],
        ["Mode", "Most frequent value", "In [1,1,2,3,3,3,4] the mode is 3"],
        ["Range", "R = max - min", "In [2,5,8,12]: R = 12-2 = 10"],
        ["Variance", "s^2 = (1/(n-1)) * Sum(x_i - x_bar)^2", "Mean squared deviation from mean"],
        ["Std. Deviation", "s = sqrt(s^2)", "Square root of variance. Same unit as data"],
        ["Coeff. of Variation", "CV = s / x_bar * 100", "Relative spread in %. Ratio scale only"],
        ["Standardization", "z = (x - x_bar) / s", "z-score: deviation in SD units. Mean 0, SD 1"],
        ["Correlation (r)", "r = cov(x,y) / (s_x * s_y)", "Range: -1 to +1. 0 = no linear relationship"],
        ["Proportion", "p = k / n", "Proportion of k successes in n trials"],
        ["Probability", "P(A) = |A| / |Omega|", "Laplace: favorable / possible outcomes"],
        ["Conditional Prob.", "P(A|B) = P(A and B) / P(B)", "Probability of A given B"],
        ["Independence", "P(A and B) = P(A) * P(B)", "If A and B are independent"],
        ["Expected Value", "E(X) = Sum(x_i * p_i)", "Mean of the random variable X"],
        ["Binomial Dist.", "P(X=k) = C(n,k) * p^k * (1-p)^(n-k)", "n trials, success probability p"],
        ["Normal Dist.", "phi(z) = 1/sqrt(2*pi) * e^(-z^2/2)", "Bell curve. Standard normal N(0,1)"],
        ["68-95-99.7 Rule", "mu +/- 1s: 68%; +/- 2s: 95%; +/- 3s: 99.7%", "Rule of thumb for normal distribution"],
        ["CI (Mean)", "CI = x_bar +/- z * s / sqrt(n)", "e.g. z=1.96 for 95% CI. Larger n = narrower CI"],
        ["Significance Level", "alpha = 0.05 (5%)", "Probability of Type I error"],
        ["p-value", "p < alpha -> reject H0", "Probability of observed result under H0"],
        ["Type I Error", "Reject H0 when H0 is true", "False positive. Risk = alpha"],
        ["Type II Error", "Retain H0 when H1 is true", "False negative. Risk = beta"],
    ],
    [3.5, 8.0, 14.0]
)

make_sheet("cheatsheet_bachelor.docx",
    "Statistics Formula Sheet — Bachelor / Standard Level",
    "Inferential Statistics, t-Test, ANOVA, Regression, Chi-Square",
    [
        ["Topic", "Formula / Concept", "Explanation / Example"],
        ["One-Sample t", "t = (x_bar - mu_0) / (s / sqrt(n))", "Compare sample mean to population mean. df = n-1"],
        ["Independent t", "t = (x1 - x2) / sqrt(s_p^2 * (1/n1 + 1/n2))", "Compare two independent groups. s_p = pooled variance"],
        ["Paired t", "t = d_bar / (s_d / sqrt(n))", "Difference per person. df = n-1"],
        ["Cohen's d", "d = (x1 - x2) / s_pooled", "Effect size: 0.2=small, 0.5=medium, 0.8=large"],
        ["F-Statistic (ANOVA)", "F = MS_between / MS_within", "Variance between / variance within groups"],
        ["Post-hoc Bonferroni", "alpha_corr = alpha / k", "k = number of comparisons. Conservative"],
        ["Tukey HSD", "HSD = q * sqrt(MS_within / n)", "Honestly Significant Difference"],
        ["Eta-squared", "eta^2 = SS_between / SS_total", "Proportion of variance explained by groups"],
        ["Simple Regression", "y = b0 + b1*x + e", "b1 = r * (s_y/s_x); b0 = y_bar - b1*x_bar"],
        ["R-squared", "R^2 = SS_regression / SS_total", "Proportion of explained variance. 0 to 1"],
        ["SE of Regression", "SE = sqrt(Sum(e_i^2) / (n-2))", "Standard error of residuals"],
        ["Correlation", "r = Sum((x-x_bar)*(y-y_bar)) / ((n-1)*s_x*s_y)", "Pearson correlation. -1 <= r <= +1"],
        ["Coeff. of Determination", "r^2", "Proportion of shared variance"],
        ["Chi-square", "chi^2 = Sum((O-E)^2 / E)", "O = observed, E = expected. Test of independence"],
        ["Cramer's V", "V = sqrt(chi^2 / (n * (min(r,c)-1)))", "Effect size for chi^2: 0.1, 0.3, 0.5"],
        ["Fisher's Exact Test", "hypergeometric probability", "For 2x2 tables with cells < 5"],
        ["CI Mean", "CI = x_bar +/- t_(n-1, alpha/2) * s / sqrt(n)", "Look up t-value from t-table"],
        ["CI Proportion", "CI = p +/- z * sqrt(p*(1-p)/n)", "e.g. z=1.96 for 95%"],
        ["Sample Size", "n = (z^2 * p * (1-p)) / e^2", "e = desired margin of error (e.g., 0.03)"],
        ["Standard Error", "SE = s / sqrt(n)", "Spread of the sampling distribution"],
        ["Effect size (chi2)", "w = sqrt(chi^2 / N)", "Cohen: 0.1=small, 0.3=medium, 0.5=large"],
    ],
    [3.5, 8.0, 14.0]
)

make_sheet("cheatsheet_master.docx",
    "Statistics Formula Sheet — Master / Advanced Level",
    "Multiple Regression, Logistic Regression, ANCOVA, Factor Analysis, Nonparametric",
    [
        ["Topic", "Formula / Concept", "Explanation / Example"],
        ["Multiple Regression", "y = b0 + b1*x1 + ... + bk*xk + e", "Multiple IVs simultaneously. Partial regression coeff."],
        ["Adjusted R^2", "R^2_adj = 1 - (1-R^2)*(n-1)/(n-k-1)", "Penalty for many predictors. Model comparison"],
        ["VIF", "VIF = 1 / (1 - R_j^2)", "R_j^2 from regressing X_j on others. VIF > 10 = trouble"],
        ["Partial Correlation", "r_xy.z = (r_xy - r_xz*r_yz) / sqrt((1-r_xz^2)*(1-r_yz^2))", "Correlation of x and y after partialling out z"],
        ["Logit (Log. Reg.)", "ln(p/(1-p)) = b0 + b1*x1 + ... + bk*xk", "Log-odds as linear function of predictors"],
        ["Odds Ratio", "OR = exp(b)", "Factor change in odds per X unit. OR > 1: increasing"],
        ["Pseudo-R^2 (Nagelkerke)", "R^2_N = (1 - (L0/Lm)^(2/n)) / (1 - L0^(2/n))", "Approximation of variance explained"],
        ["AIC", "AIC = -2*ln(L) + 2*k", "k = number of parameters. Lower AIC = better model"],
        ["BIC", "BIC = -2*ln(L) + k*ln(n)", "More strict than AIC (favors simpler models)"],
        ["Partial Eta^2", "eta^2_p = SS_effect / (SS_effect + SS_error)", "Effect size per effect in ANOVA/ANCOVA"],
        ["Eigenvalue (FA)", "Eigenvalue = Sum(ai^2) for factor i", "ai = loadings. Kaiser: eigenvalue > 1 retained"],
        ["KMO", "KMO = Sum(rij^2) / (Sum(rij^2) + Sum(aij^2))", "KMO > 0.6 = acceptable. Sampling Adequacy"],
        ["Bartlett's Test", "chi^2 = -((n-1) - (2p+5)/6)*ln(|R|)", "Tests non-diagonal structure. p<.05 = suitable"],
        ["CFI (CFA)", "CFI = 1 - max(chi^2_m - df_m, 0) / max(chi^2_m - df_m, chi^2_b - df_b, 0)", "CFI > .95 = good fit. max() prevents out-of-range values"],
        ["RMSEA", "RMSEA = sqrt(max((chi^2_m - df_m)/(N-1), 0) / df_m)", "RMSEA < .06 = good; < .08 acceptable"],
        ["Mann-Whitney U", "U = Sum(R1) - n1*(n1+1)/2", "Rank sum group 1 minus minimum possible"],
        ["Kruskal-Wallis H", "H = 12/(N(N+1)) * Sum(R_i^2/n_i) - 3(N+1)", "Rank-based ANOVA. ~ chi^2 with k-1 df"],
        ["Friedman", "chi^2_F = 12n/(k(k+1)) * Sum(R_j^2) - 3n(k+1)", "Rank-based RM-ANOVA. R_j = sum of ranks per time"],
    ],
    [3.5, 8.0, 14.0]
)

make_sheet("cheatsheet_phd.docx",
    "Statistics Formula Sheet — PhD / All-Access Level",
    "Bayesian Statistics, Multilevel, SEM, Machine Learning, Meta-Analysis",
    [
        ["Topic", "Formula / Concept", "Explanation / Example"],
        ["Bayes' Theorem", "P(theta|D) = P(D|theta)*P(theta) / P(D)", "Posterior = Likelihood * Prior / Evidence"],
        ["Bayes Factor", "BF10 = P(D|H1) / P(D|H0)", "BF > 3: moderate; BF > 10: strong evidence"],
        ["MCMC", "P(theta|D) ~ 1/m * Sum(g(theta^(t)))", "Posterior mean via MCMC. R-hat < 1.01"],
        ["ICC (Multilevel)", "ICC = sigma^2_between / (sigma^2_between + sigma^2_within)", "Proportion of level-2 variance. ICC > .05 => HLM"],
        ["Multilevel Model", "y_ij = gamma_00 + gamma_10*x_ij + u_0j + u_1j*x_ij + e_ij", "Fixed: gamma; Random: u; Level-1: i; Level-2: j"],
        ["Cross-Level Interaction", "gamma_11 = effect of w_j (L2) on b_1j (L1 slope)", "Does a level-2 variable moderate a level-1 effect?"],
        ["SEM - CFI", "CFI = 1 - max(chi^2_m - df_m, 0) / max(chi^2_m - df_m, chi^2_b - df_b, 0)", "> .95: good; > .90: acceptable"],
        ["SEM - RMSEA", "RMSEA = sqrt(max(chi^2_m - df_m, 0) / (df_m * (N-1)))", "< .06: good; .06-.08: acceptable; > .10: poor"],
        ["Mediation (PROCESS)", "a*b = indirect effect; c' = direct effect", "X -> M (a); M -> Y (b); a*b boot CI != 0 = sig. mediation"],
        ["Moderation", "y = b0 + b1*x + b2*z + b3*x*z + e", "b3 = interaction. Simple slopes at Mean +/- 1 SD of z"],
        ["Power Analysis", "n = f(alpha, beta, d)", "Power > .80. Use G*Power. Effect size from prior work"],
        ["Meta (Fixed Effects)", "theta_pooled = Sum(w_i * theta_i) / Sum(w_i)", "w_i = 1/v_i (inverse variance). One true effect"],
        ["Meta (Random Effects)", "w_i^* = 1 / (v_i + tau^2)", "tau^2 = between-study variance. Effect varies"],
        ["I^2", "I^2 = (Q - df) / Q * 100%", "Heterogeneity proportion. 25/50/75% = low/medium/high"],
        ["Funnel Plot", "Effect (x) vs. SE (y)", "Asymmetry = publication bias. Egger test p < .10"],
        ["Precision-Recall", "P = TP/(TP+FP); R = TP/(TP+FN)", "Precision: correctness. Recall: completeness"],
        ["F1-Score", "F1 = 2 * P * R / (P + R)", "Harmonic mean of precision and recall"],
        ["AUC-ROC", "AUC = Integral(TPR(FPR)) d(FPR)", "AUC > .70 acceptable; > .80 good; > .90 excellent"],
        ["Bias-Variance", "E[(y - f(x))^2] = Bias^2 + Var + sigma^2", "Bias: systematic error; Variance: spread. Tradeoff"],
    ],
    [3.5, 8.0, 14.0]
)

print("\nEnglische Spickzettel erfolgreich erstellt.")
